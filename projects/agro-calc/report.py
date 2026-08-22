# -*- coding: utf-8 -*-
"""Формирование отчёта (.docx) по расчёту кормления или посева."""
import datetime

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

from norms import ANIMALS, FEED_TYPES, CROPS, FEEDING_REGIMES

DISCLAIMER = (
    "Расчёт выполнен по типовым ориентировочным нормам (справочные "
    "зоотехнические/агрономические коэффициенты). Точные значения зависят "
    "от породы/сорта, зоны, погоды и качества кормов/почвы — перед "
    "применением в хозяйстве уточните нормы у зоотехника/агронома."
)


def _base_style(doc):
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)


def _h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _p(doc, text="", bold=False, italic=False, size=12, align=None):
    par = doc.add_paragraph()
    if align is not None:
        par.alignment = align
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Times New Roman"
    return par


def _table(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        for run in hdr[i].paragraphs[0].runs:
            run.bold = True
    for row in rows:
        cells = t.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
    return t


def _round(v, nd=1):
    try:
        return round(v, nd)
    except TypeError:
        return v


def _header(doc, title, author, rdate):
    _p(doc, "Агро-калькулятор — " + title, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    if author:
        _p(doc, f"Составил: {author}", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    _p(doc, f"Дата: {rdate or datetime.date.today().isoformat()}", align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    doc.add_paragraph()


def generate_feed_report(ctx, out_path):
    doc = Document()
    _base_style(doc)
    animal = ANIMALS[ctx["animal_id"]]
    _header(doc, f"кормление: {animal['label']}", ctx.get("author"), ctx.get("date"))

    data = ctx["result"]
    per_head = data["per_head"]

    _h1(doc, "1. Исходные данные")
    _table(doc, ["Параметр", "Значение"], [
        ["Животное", animal["label"]],
        ["Живая масса, кг", data["weight_kg"]],
        ["Продуктивность на голову", f"{data['output_per_head']} ({animal['production']['unit_label']})"],
        ["Поголовье", data["count"]],
        ["Период, суток", data["days"]],
        ["Режим кормления", FEEDING_REGIMES.get(data["regime"], data["regime"])],
    ])
    doc.add_paragraph()

    _h1(doc, "2. Суточная потребность на 1 голову")
    _table(doc, ["Показатель", "Значение"], [
        ["ЭКЕ на поддержание", f"{_round(per_head['maintenance_efu'])} ЭКЕ/сут"],
        ["ЭКЕ на продукцию", f"{_round(per_head['production_efu'])} ЭКЕ/сут"],
        ["Итого ЭКЕ/сутки на голову", f"{_round(per_head['total_efu_per_head'])} ЭКЕ/сут"],
    ])
    doc.add_paragraph()

    _h1(doc, "3. Раскладка рациона на 1 голову в сутки")
    rows = [
        [FEED_TYPES[f]["label"], _round(kg, 2)]
        for f, kg in per_head["feed_kg_per_head"].items() if kg > 0.001
    ]
    _table(doc, ["Вид корма", "кг/сутки на голову"], rows)
    doc.add_paragraph()

    _h1(doc, f"4. Итого на всё поголовье ({data['count']} гол.) за {data['days']} суток")
    rows = [
        [FEED_TYPES[f]["label"], _round(kg, 1)]
        for f, kg in data["feed_kg_period_all"].items() if kg > 0.01
    ]
    _table(doc, ["Вид корма", "кг за период"], rows)
    _p(doc, f"Всего ЭКЕ за период на всё поголовье: {_round(data['total_efu_period_all'])} ЭКЕ", bold=True)
    doc.add_paragraph()

    if "required_count" in data:
        _h1(doc, "5. Обратный расчёт — по желаемому результату")
        _p(doc, f"Желаемый суммарный результат за период: {data['target_total_output']}")
        _p(doc, f"Требуемое поголовье: {data['required_count']} гол.", bold=True)
        doc.add_paragraph()

    doc.add_paragraph()
    _p(doc, DISCLAIMER, italic=True, size=10)

    doc.save(out_path)


def generate_crop_report(ctx, out_path):
    doc = Document()
    _base_style(doc)
    crop = CROPS[ctx["crop_id"]]
    _header(doc, f"посев: {crop['label']}", ctx.get("author"), ctx.get("date"))

    data = ctx["result"]

    _h1(doc, "1. Исходные данные")
    rows = [
        ["Культура", crop["label"]],
        ["Площадь, га", _round(data["area_ha"], 3)],
        ["Урожайность, т/га", data["yield_t_ha"]],
    ]
    if data.get("sowing_date"):
        rows.append(["Дата посева", data["sowing_date"]])
    _table(doc, ["Параметр", "Значение"], rows)
    doc.add_paragraph()

    if "required_area_ha" in data:
        _h1(doc, "2. Обратный расчёт — по желаемому сбору урожая")
        _p(doc, f"Желаемый сбор: {data['target_total_yield_t']} т")
        _p(doc, f"Требуемая площадь: {_round(data['required_area_ha'], 3)} га", bold=True)
        doc.add_paragraph()

    _h1(doc, "3. Семена и ожидаемый сбор")
    _table(doc, ["Показатель", "Значение"], [
        ["Норма высева, кг/га", crop["seeding_rate_kg_ha"]],
        ["Нужно семян всего, кг", _round(data["seed_kg"], 1)],
        ["Ожидаемый сбор, т", _round(data["total_yield_t"], 2)],
    ])
    doc.add_paragraph()

    _h1(doc, "4. Удобрения (по выносу питательных веществ с урожаем)")
    rows = [[n, _round(kg, 1)] for n, kg in data["nutrient_kg"].items()]
    _table(doc, ["Элемент", "кг действующего вещества"], rows)
    doc.add_paragraph()
    if data["fertilizer_products"]:
        rows = [
            [info["fertilizer_label"], _round(info["product_kg"], 1)]
            for info in data["fertilizer_products"].values()
        ]
        _table(doc, ["Удобрение", "кг товарного продукта"], rows)
    doc.add_paragraph()

    _h1(doc, "5. Вода и трудозатраты")
    _table(doc, ["Показатель", "Значение"], [
        ["Суммарная потребность культуры в воде за сезон (осадки + полив), м³", _round(data["water_m3"], 0)],
        ["Трудозатраты, чел.-дней", _round(data["labor_days"], 1)],
    ])
    _p(doc, "Если часть влаги обеспечивают осадки — вычтите их из этого объёма, "
             "чтобы получить объём полива.", italic=True, size=10)
    doc.add_paragraph()

    if data.get("harvest_range"):
        _h1(doc, "6. Календарь созревания")
        _p(doc, f"Посев: {data['sowing_date']}")
        _p(doc, f"Ожидаемая уборка: {data['harvest_range'][0]} — {data['harvest_range'][1]}", bold=True)
        doc.add_paragraph()

    if ctx.get("econ"):
        econ = ctx["econ"]
        _h1(doc, "7. Экономика")
        rows = [
            ["Выручка", _round(econ["revenue"])],
            ["Затраты на семена", _round(econ["seed_cost"])],
            ["Затраты на удобрения", _round(econ["fert_cost"])],
            ["Затраты на воду", _round(econ["water_cost"])],
            ["Затраты на оплату труда", _round(econ["labor_cost"])],
            ["Итого затрат", _round(econ["total_cost"])],
            ["Прибыль", _round(econ["profit"])],
        ]
        _table(doc, ["Статья", "Сумма"], rows)
        doc.add_paragraph()

    doc.add_paragraph()
    _p(doc, DISCLAIMER, italic=True, size=10)

    doc.save(out_path)

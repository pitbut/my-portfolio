# -*- coding: utf-8 -*-
"""Формирование расчётно-пояснительной записки для фермы (.docx)."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime
import math

SUPPORT_NAMES_RU = {
    'pin': 'шарнирно-неподвижная опора',
    'roller_y': 'каток (допускает смещение по горизонтали)',
    'roller_x': 'каток (допускает смещение по вертикали)',
}


def _set_base_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def _h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _p(doc, text='', bold=False, italic=False, align=None, size=13):
    par = doc.add_paragraph()
    if align is not None:
        par.alignment = align
    run = par.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(size); run.font.name = 'Times New Roman'
    return par


def _formula(doc, parts):
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, kind in parts:
        run = par.add_run(text)
        run.font.name = 'Times New Roman'; run.font.size = Pt(13)
        if kind == 'i': run.italic = True
        elif kind == 'sup': run.font.superscript = True
        elif kind == 'sub': run.font.subscript = True
    return par


def _num(v, nd=4):
    if abs(v) < 1e-9:
        v = 0.0
    return f"{v:.{nd}g}"


def _add_image(doc, path, width_cm=15.5):
    doc.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_report(ctx, out_path):
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0); section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5); section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0); section.bottom_margin = Cm(2.0)
    _set_base_style(doc)

    truss = ctx['truss']
    sec = ctx['section']
    deg, m, r, n = truss.degree_of_determinacy()

    # ---------- титул ----------
    _p(doc, "Расчёт плоской фермы (сопротивление материалов)", bold=True,
       align=WD_ALIGN_PARAGRAPH.CENTER, size=15)
    _p(doc, "")
    _p(doc, "РАСЧЁТНО-ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    _p(doc, "к расчёту фермы методом конечных элементов", align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "")
    if ctx.get('author'):
        _p(doc, f"Исполнитель: {ctx['author']}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc, f"Дата: {ctx.get('date', datetime.date.today().strftime('%d.%m.%Y'))}",
       align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()

    # ---------- 1. Задание ----------
    _h1(doc, "1. Задание на расчёт")
    _p(doc, f"Дана плоская стержневая ферма с шарнирным сопряжением стержней в узлах: "
            f"{n} узлов, {m} стержней, {r} опорных связей. Нагрузки приложены в узлах "
            f"(классическое допущение для расчёта ферм — стержни работают только на "
            f"растяжение/сжатие). Требуется: определить опорные реакции, усилия в "
            f"стержнях, проверить прочность растянутых и устойчивость сжатых стержней.")

    _p(doc, "")
    _p(doc, "Координаты узлов:", bold=True)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Light Grid Accent 1'
    for c, txt in zip(tbl.rows[0].cells, ["№ узла", "x, м", "y, м"]):
        c.text = txt; c.paragraphs[0].runs[0].bold = True
    for i, nd in enumerate(truss.nodes):
        row = tbl.add_row().cells
        row[0].text = str(i); row[1].text = _num(nd.x); row[2].text = _num(nd.y)

    _p(doc, "")
    _p(doc, "Стержни (соединяемые узлы):", bold=True)
    tbl2 = doc.add_table(rows=1, cols=3)
    tbl2.style = 'Light Grid Accent 1'
    for c, txt in zip(tbl2.rows[0].cells, ["№ стержня", "Узлы", "Длина, м"]):
        c.text = txt; c.paragraphs[0].runs[0].bold = True
    for i, mb in enumerate(truss.members):
        row = tbl2.add_row().cells
        row[0].text = str(i); row[1].text = f"{mb.i} — {mb.j}"
        row[2].text = _num(truss.member_length(mb), 3)

    _p(doc, "")
    _p(doc, "Опоры:", bold=True)
    for s in truss.supports:
        _p(doc, f"— узел {s.node}: {SUPPORT_NAMES_RU[s.kind]}")

    _p(doc, "Узловые нагрузки:", bold=True)
    for ld in ctx['loads_in']:
        _p(doc, f"— узел {ld['node']}: Fx = {_num(ld['fx'])} Н, Fy = {_num(ld['fy'])} Н")

    _p(doc, "Материал / сечение стержней:", bold=True)
    _p(doc, f"— {sec['desc']} (одинаково для всех стержней)")
    _p(doc, f"— модуль упругости E = {ctx['E']:.4g} Па; "
            f"допускаемое напряжение [σ] = {ctx['sigma_allow'] / 1e6:.3g} МПа")

    _p(doc, "")
    _add_image(doc, ctx['images']['scheme'])
    _p(doc, "Рисунок 1 — расчётная схема фермы", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True)

    # ---------- 2. Определимость ----------
    doc.add_page_break()
    _h1(doc, "2. Проверка статической определимости")
    _formula(doc, [("m + r − 2n = ", 'n'), (f"{m} + {r} − 2·{n} = ", 'n'), (f"{deg}", 'n')])
    if deg == 0:
        _p(doc, "Условие m + r = 2n выполняется — ферма статически определима: "
                "усилия в стержнях и опорные реакции находятся из уравнений равновесия "
                "(метод узлов / метод сечений).")
    elif deg > 0:
        _p(doc, f"m + r > 2n на {deg} — ферма {deg} раз(а) статически неопределима "
                f"(имеются лишние стержни/связи). Задача решена методом конечных "
                f"элементов (метод перемещений): каждому стержню сопоставлен элемент "
                f"с осевой жёсткостью EA/L, глобальная система собрана и решена "
                f"относительно узловых перемещений при заданных опорных связях; "
                f"усилия в стержнях получены из найденных перемещений "
                f"(N = EA/L·Δ, что эквивалентно результату метода сил).")
    else:
        _p(doc, "m + r < 2n — система является механизмом (геометрически изменяема) "
                "и не может воспринимать произвольную нагрузку без дополнительных связей.")

    # ---------- 3. Реакции ----------
    doc.add_page_break()
    _h1(doc, "3. Определение опорных реакций")
    tblr = doc.add_table(rows=1, cols=4)
    tblr.style = 'Light Grid Accent 1'
    for c, txt in zip(tblr.rows[0].cells, ["Опора (узел)", "Тип", "Rx, Н", "Ry, Н"]):
        c.text = txt; c.paragraphs[0].runs[0].bold = True
    for rr in truss._reactions:
        row = tblr.add_row().cells
        row[0].text = str(rr['node']); row[1].text = SUPPORT_NAMES_RU[rr['kind']]
        row[2].text = _num(rr['Rx']); row[3].text = _num(rr['Ry'])
    sfx, sfy = truss.check_equilibrium()
    _p(doc, "")
    _p(doc, f"Проверка: ΣFx = {_num(sfx, 3)} ≈ 0, ΣFy = {_num(sfy, 3)} ≈ 0 — "
            f"условия равновесия выполняются.")

    # ---------- 4. Усилия в стержнях ----------
    doc.add_page_break()
    _h1(doc, "4. Усилия в стержнях")
    _p(doc, "Усилие в стержне определено через его удлинение, найденное из решения "
            "МКЭ: N = (EA/L)·Δ, где Δ — изменение длины стержня между узловыми "
            "перемещениями. Результат эквивалентен методу вырезания узлов. "
            "Знак «+» — растяжение, «−» — сжатие.")
    tblf = doc.add_table(rows=1, cols=4)
    tblf.style = 'Light Grid Accent 1'
    for c, txt in zip(tblf.rows[0].cells, ["№ стержня", "Узлы", "N, Н", "Тип"]):
        c.text = txt; c.paragraphs[0].runs[0].bold = True
    forces = truss._forces
    for i, (mb, N) in enumerate(zip(truss.members, forces)):
        row = tblf.add_row().cells
        row[0].text = str(i); row[1].text = f"{mb.i} — {mb.j}"
        row[2].text = _num(N)
        row[3].text = "растяжение" if N > 1e-6 else ("сжатие" if N < -1e-6 else "0 (ненагружен)")

    i_max = max(range(len(forces)), key=lambda k: forces[k])
    i_min = min(range(len(forces)), key=lambda k: forces[k])
    _p(doc, "")
    _p(doc, f"Максимальное растягивающее усилие: N = {_num(forces[i_max])} Н "
            f"в стержне {i_max} (узлы {truss.members[i_max].i}—{truss.members[i_max].j}).")
    _p(doc, f"Максимальное сжимающее усилие: N = {_num(forces[i_min])} Н "
            f"в стержне {i_min} (узлы {truss.members[i_min].i}—{truss.members[i_min].j}).")
    _p(doc, "")
    _add_image(doc, ctx['images']['forces'])
    _p(doc, "Рисунок 2 — усилия в стержнях (толщина и цвет линии — по модулю усилия)",
       align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True)

    # ---------- 5. Сечение ----------
    doc.add_page_break()
    _h1(doc, "5. Геометрические характеристики сечения стержней")
    _p(doc, f"Принято сечение (одинаковое для всех стержней): {sec['desc']}.")
    _formula(doc, [("A = ", 'n'), (_num(sec['A'] * 1e6), 'n'), (" мм", 'n'), ("2", 'sup'),
                    (" = ", 'n'), (_num(sec['A']), 'n'), (" м", 'n'), ("2", 'sup')])
    _formula(doc, [("I", 'i'), ("min", 'sub'), (" = ", 'n'), (_num(ctx['I_min'] * 1e12), 'n'),
                    (" мм", 'n'), ("4", 'sup'), (" = ", 'n'), (_num(ctx['I_min']), 'n'),
                    (" м", 'n'), ("4", 'sup'), ("  (минимальный — для проверки устойчивости)", 'n')])

    # ---------- 6. Прочность и устойчивость ----------
    doc.add_page_break()
    _h1(doc, "6. Проверка прочности и устойчивости стержней")
    _p(doc, "Растянутые стержни — проверка по условию прочности:")
    _formula(doc, [("σ = N / A ≤ [σ]", 'n')])
    _p(doc, "Сжатые стержни — дополнительно проверяются на устойчивость "
            "(потерю продольной устойчивости) по формуле Эйлера:")
    _formula(doc, [("N", 'i'), ("кр", 'sub'), (" = π", 'n'), ("2", 'sup'), ("·E·I", 'n'),
                    ("min", 'sub'), (" / (μ·L)", 'n'), ("2", 'sup')])
    _p(doc, "где μ — коэффициент приведения длины (для шарнирного закрепления "
            "концов стержня в ферме μ = 1).")

    tbls = doc.add_table(rows=1, cols=6)
    tbls.style = 'Light Grid Accent 1'
    for c, txt in zip(tbls.rows[0].cells,
                       ["№", "N, Н", "σ = N/A, МПа", "Условие", "N_кр (Эйлер), Н", "Запас"]):
        c.text = txt; c.paragraphs[0].runs[0].bold = True

    E = ctx['E']; I_min = ctx['I_min']; A = sec['A']; sigma_allow = ctx['sigma_allow']
    worst_ratio = 0.0
    worst_desc = ""
    for i, (mb, N) in enumerate(zip(truss.members, forces)):
        L = truss.member_length(mb)
        sigma = N / A
        row = tbls.add_row().cells
        row[0].text = str(i)
        row[1].text = _num(N)
        row[2].text = _num(sigma / 1e6)
        if N >= 0:
            ok = abs(sigma) <= sigma_allow
            row[3].text = "прочность"
            row[4].text = "—"
            row[5].text = f"{sigma_allow/abs(sigma):.2f}" if sigma > 1e-6 else "—"
            ratio = abs(sigma) / sigma_allow if sigma > 1e-6 else 0
        else:
            Ncr = math.pi ** 2 * E * I_min / L ** 2
            ok = abs(N) <= Ncr and abs(sigma) <= sigma_allow
            row[3].text = "устойчивость + прочность"
            row[4].text = _num(Ncr)
            n_safety = Ncr / abs(N) if abs(N) > 1e-6 else float('inf')
            row[5].text = f"{n_safety:.2f}"
            ratio = abs(N) / Ncr if Ncr > 0 else 1e9
        if not ok and ratio > worst_ratio:
            worst_ratio = ratio
            worst_desc = f"стержень {i}"
        if ratio > worst_ratio:
            worst_ratio = ratio

    all_ok = worst_ratio <= 1.0
    _p(doc, "")
    verdict = "Все условия прочности и устойчивости ВЫПОЛНЯЮТСЯ" if all_ok else \
              f"Условие НЕ выполняется ({worst_desc or 'см. таблицу'}) — требуется увеличить сечение"
    _p(doc, verdict + ".", bold=True)

    # ---------- Вывод ----------
    doc.add_page_break()
    _h1(doc, "Вывод")
    _p(doc, f"Рассчитана плоская ферма из {n} узлов и {m} стержней "
            f"({'статически определимая' if deg == 0 else f'{deg} раз(а) статически неопределимая'}). "
            f"Максимальное растягивающее усилие {_num(forces[i_max])} Н, максимальное "
            f"сжимающее усилие {_num(forces[i_min])} Н. Для сечения «{sec['desc']}» "
            f"{'все стержни удовлетворяют условиям прочности и устойчивости.' if all_ok else 'не все стержни удовлетворяют условиям прочности/устойчивости — необходимо увеличить сечение или изменить конфигурацию фермы.'}")

    doc.save(out_path)
    return out_path

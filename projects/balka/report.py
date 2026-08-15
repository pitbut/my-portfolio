# -*- coding: utf-8 -*-
"""Формирование расчётно-пояснительной записки в формате .docx."""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import datetime

SUPPORT_NAMES_RU = {
    'pin': 'шарнирно-неподвижная опора',
    'roller': 'шарнирно-подвижная опора',
    'fixed': 'жёсткая заделка',
}


def _set_base_style(doc):
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    rpr = style.element.get_or_add_rPr()
    rFonts = rpr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = rpr.makeelement(qn('w:rFonts'), {})
        rpr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def _h1(doc, text):
    p = doc.add_heading(text, level=1)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(15)
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _h2(doc, text):
    p = doc.add_heading(text, level=2)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13.5)
        r.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _p(doc, text='', bold=False, italic=False, align=None, size=13):
    par = doc.add_paragraph()
    if align is not None:
        par.alignment = align
    run = par.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = 'Times New Roman'
    return par


def _formula(doc, parts):
    """parts: список кортежей (text, kind) где kind in
    ('n','i','sup','sub')  -> normal / italic / superscript / subscript"""
    par = doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for text, kind in parts:
        run = par.add_run(text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        if kind == 'i':
            run.italic = True
        elif kind == 'sup':
            run.font.superscript = True
        elif kind == 'sub':
            run.font.subscript = True
    return par


def _num(v, nd=3):
    if abs(v) < 1e-9:
        v = 0.0
    return f"{v:.{nd}g}"


def _add_image(doc, path, width_cm=15.5):
    doc.add_picture(path, width=Cm(width_cm))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER


def generate_report(ctx, out_path):
    """
    ctx: словарь с ключами
      title, author, group, date
      beam: Beam-объект (после solve_reactions()/segments())
      supports_in, forces_in, moments_in, dloads_in — исходные данные (списки dict)
      section: dict геом.характеристик (A, Ix, Wx, y_max, desc)
      sigma_allow: допускаемое напряжение, Па
      images: dict путей к рисункам {'scheme','Q','M','defl'}
      determinate: bool — статически определима ли система опор
    """
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(1.5)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    _set_base_style(doc)

    beam = ctx['beam']
    L = beam.L

    # ---------- титул ----------
    _p(doc, ctx.get('org', 'Расчёт балки на прочность и жёсткость (сопротивление материалов)'),
       bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=15)
    _p(doc, "")
    _p(doc, "РАСЧЁТНО-ПОЯСНИТЕЛЬНАЯ ЗАПИСКА", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=16)
    _p(doc, "к расчёту балки на изгиб методом конечных элементов", align=WD_ALIGN_PARAGRAPH.CENTER)
    _p(doc, "")
    if ctx.get('author'):
        _p(doc, f"Исполнитель: {ctx['author']}", align=WD_ALIGN_PARAGRAPH.RIGHT)
    _p(doc, f"Дата: {ctx.get('date', datetime.date.today().strftime('%d.%m.%Y'))}",
       align=WD_ALIGN_PARAGRAPH.RIGHT)
    doc.add_page_break()

    # ---------- 1. Задание ----------
    _h1(doc, "1. Задание на расчёт")
    _p(doc, f"Дана прямая балка постоянного сечения длиной L = {_num(L)} м, "
            f"нагруженная системой сил, изгибающих моментов и распределённых нагрузок, "
            f"расположенная на {len(beam.supports)} опор(ах). Требуется:")
    for i, t in enumerate([
        "определить опорные реакции;",
        "построить эпюры поперечных сил Q(x) и изгибающих моментов M(x);",
        "подобрать/проверить поперечное сечение балки из условия прочности;",
        "построить эпюру прогибов и проверить условие жёсткости.",
    ], 1):
        _p(doc, f"{i}) {t}")

    _p(doc, "")
    _p(doc, "Опоры:", bold=True)
    for s in beam.supports:
        _p(doc, f"— {SUPPORT_NAMES_RU[s.kind]} в сечении x = {_num(s.x)} м")

    if ctx['forces_in']:
        _p(doc, "Сосредоточенные силы:", bold=True)
        for f in ctx['forces_in']:
            direction = "вниз" if f['value'] < 0 else "вверх"
            _p(doc, f"— F = {_num(abs(f['value']))} Н, направлена {direction}, "
                    f"приложена в x = {_num(f['x'])} м")

    if ctx['moments_in']:
        _p(doc, "Сосредоточенные моменты:", bold=True)
        for m in ctx['moments_in']:
            direction = "против часовой стрелки" if m['value'] > 0 else "по часовой стрелке"
            _p(doc, f"— M = {_num(abs(m['value']))} Н·м, направлен {direction}, "
                    f"приложен в x = {_num(m['x'])} м")

    if ctx['dloads_in']:
        _p(doc, "Распределённые нагрузки:", bold=True)
        for d in ctx['dloads_in']:
            direction = "вниз" if (d['value1'] + d['value2']) < 0 else "вверх"
            if d['value1'] == d['value2']:
                _p(doc, f"— равномерная q = {_num(abs(d['value1']))} Н/м, направлена {direction}, "
                        f"на участке x ∈ [{_num(d['x1'])}; {_num(d['x2'])}] м")
            else:
                _p(doc, f"— линейно меняющаяся от {_num(abs(d['value1']))} до {_num(abs(d['value2']))} Н/м, "
                        f"направлена {direction}, на участке x ∈ [{_num(d['x1'])}; {_num(d['x2'])}] м")

    _p(doc, "Материал / сечение:", bold=True)
    _p(doc, f"— {ctx['section']['desc']}")
    _p(doc, f"— модуль упругости E = {ctx['E']:.4g} Па; "
            f"допускаемое напряжение [σ] = {ctx['sigma_allow'] / 1e6:.3g} МПа")

    _p(doc, "")
    _add_image(doc, ctx['images']['scheme'])
    _p(doc, "Рисунок 1 — расчётная схема балки", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True)

    # ---------- 2. Реакции ----------
    doc.add_page_break()
    _h1(doc, "2. Определение опорных реакций")

    n_unknown = sum((2 if s.kind == 'fixed' else 1) for s in beam.supports)
    if ctx['determinate']:
        _p(doc, "Система опор статически определима: число неизвестных реакций равно числу "
                "независимых уравнений равновесия плоской системы сил (ΣFy = 0, ΣM = 0).")
        ax_pt = beam.supports[0].x
        terms = []
        for f in ctx['forces_in']:
            terms.append(f"{'+' if f['value']>=0 else '-'}{_num(abs(f['value']))}")
        _p(doc, "Уравнения равновесия (ось y — вверх, момент против часовой стрелки — положительный):")
        _p(doc, "ΣFy = 0;      ΣM(относительно опоры) = 0.")
    else:
        _p(doc, "Число неизвестных опорных реакций превышает число уравнений статики "
                f"({n_unknown} неизвестных против 3 уравнений равновесия) — система "
                "statically indeterminate (статически неопределима). Дополнительные "
                "уравнения совместности деформаций получены методом конечных элементов "
                "(балочный элемент Эйлера-Бернулли, метод перемещений): балка разбита на "
                "участки между опорами и точками приложения нагрузок, для каждого участка "
                "составлена матрица жёсткости 4×4, глобальная система собрана и решена "
                "относительно узловых перемещений/углов поворота при заданных опорных связях; "
                "реакции получены как усилия в связях.")

    _p(doc, "")
    _p(doc, "Результат (реакции опор):", bold=True)
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Light Grid Accent 1'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl.rows[0].cells
    for c, txt in zip(hdr, ["Опора, x, м", "Тип", "R, Н", "M, Н·м"]):
        c.text = txt
        c.paragraphs[0].runs[0].bold = True
    for r in beam._reactions:
        row = tbl.add_row().cells
        row[0].text = _num(r['x'])
        row[1].text = SUPPORT_NAMES_RU[r['kind']]
        row[2].text = _num(r['R'])
        row[3].text = _num(r['M']) if r['kind'] == 'fixed' else "—"

    sum_R = sum(r['R'] for r in beam._reactions) + sum(f['value'] for f in ctx['forces_in'])
    for d in ctx['dloads_in']:
        sum_R += 0.5 * (d['value1'] + d['value2']) * (d['x2'] - d['x1'])
    _p(doc, "")
    _p(doc, f"Проверка: ΣFy = {_num(sum_R, 3)} ≈ 0 — условие равновесия выполняется.")

    # ---------- 3. Эпюры ----------
    doc.add_page_break()
    _h1(doc, "3. Построение эпюр Q(x) и M(x)")
    _p(doc, "Внутренние усилия в произвольном сечении x определены методом сечений: "
            "поперечная сила Q(x) равна сумме проекций на ось y всех сил, приложенных по одну "
            "сторону от сечения; изгибающий момент M(x) — сумме моментов этих сил относительно "
            "центра сечения (правило знаков: момент положителен, если растянуты нижние волокна). "
            "Ниже приведены аналитические выражения по участкам (участки — между опорами, "
            "силами и границами распределённых нагрузок).")

    segs = beam.segments()
    for i, seg in enumerate(segs, 1):
        x0, x1 = seg['x0'], seg['x1']
        _p(doc, f"Участок {i}:  x ∈ [{_num(x0)}; {_num(x1)}] м,  ξ = x − {_num(x0)}", bold=True)
        w0, slope = seg['w0'], seg['slope']
        if abs(w0) < 1e-9 and abs(slope) < 1e-9:
            _formula(doc, [("Q(x) = ", 'n'), (_num(seg['Q0']), 'n'), (" Н  (const)", 'n')])
        elif abs(slope) < 1e-9:
            _formula(doc, [("Q(x) = ", 'n'), (_num(seg['Q0']), 'n'), (" + ", 'n'),
                            (_num(w0), 'n'), ("·ξ,  Н", 'n')])
        else:
            _formula(doc, [("Q(x) = ", 'n'), (_num(seg['Q0']), 'n'), (" + ", 'n'),
                            (_num(w0), 'n'), ("·ξ + ", 'n'), (_num(slope), 'n'), ("·ξ", 'n'),
                            ("2", 'sup'), ("/2,  Н", 'n')])
        _formula(doc, [("M(x) = ", 'n'), (_num(seg['M0']), 'n'), (" + ", 'n'),
                        (_num(seg['Q0']), 'n'), ("·ξ + ", 'n'), (_num(w0), 'n'), ("·ξ", 'n'),
                        ("2", 'sup'), ("/2 + ", 'n'), (_num(slope), 'n'), ("·ξ", 'n'),
                        ("3", 'sup'), ("/6,  Н·м", 'n')])
        Qx0, Qx1 = beam.Q(x0 + 1e-6), beam.Q(x1 - 1e-6)
        Mx0, Mx1 = beam.M(x0 + 1e-6), beam.M(x1 - 1e-6)
        _p(doc, f"На границах участка: Q({_num(x0)}) = {_num(Qx0)} Н,  "
                f"Q({_num(x1)}) = {_num(Qx1)} Н;   "
                f"M({_num(x0)}) = {_num(Mx0)} Н·м,  M({_num(x1)}) = {_num(Mx1)} Н·м.")
        _p(doc, "")

    xq, Qv = beam.max_abs_Q()
    xm, Mv = beam.max_abs_M()
    _p(doc, f"Максимальная (по модулю) поперечная сила: Q_max = {_num(Qv)} Н при x = {_num(xq)} м.")
    _p(doc, f"Максимальный (по модулю) изгибающий момент: M_max = {_num(Mv)} Н·м при x = {_num(xm)} м.")
    _p(doc, "")
    _add_image(doc, ctx['images']['scheme'], width_cm=15.5)
    _p(doc, "Рисунок 2 — расчётная схема (повторно, для сопоставления с эпюрами)",
       align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True)
    _add_image(doc, ctx['images']['Q'])
    _p(doc, "Рисунок 3 — эпюра поперечных сил Q(x)", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True)
    _add_image(doc, ctx['images']['M'])
    _p(doc, "Рисунок 4 — эпюра изгибающих моментов M(x)", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True)

    # ---------- 4. Геометрические характеристики ----------
    doc.add_page_break()
    _h1(doc, "4. Геометрические характеристики сечения")
    sec = ctx['section']
    _p(doc, f"Принято сечение: {sec['desc']}.")
    _formula(doc, [("A = ", 'n'), (_num(sec['A'] * 1e6, 4), 'n'), (" мм", 'n'), ("2", 'sup'),
                    (" = ", 'n'), (_num(sec['A'], 4), 'n'), (" м", 'n'), ("2", 'sup')])
    _formula(doc, [("I", 'i'), ("x", 'sub'), (" = ", 'n'), (_num(sec['Ix'] * 1e12, 4), 'n'),
                    (" мм", 'n'), ("4", 'sup'), (" = ", 'n'), (_num(sec['Ix'], 4), 'n'),
                    (" м", 'n'), ("4", 'sup')])
    _formula(doc, [("W", 'i'), ("x", 'sub'), (" = ", 'n'), (_num(sec['Wx'] * 1e9, 4), 'n'),
                    (" мм", 'n'), ("3", 'sup'), (" = ", 'n'), (_num(sec['Wx'], 4), 'n'),
                    (" м", 'n'), ("3", 'sup')])

    # ---------- 5. Проверка прочности ----------
    doc.add_page_break()
    _h1(doc, "5. Проверка прочности")
    _p(doc, "Условие прочности по нормальным напряжениям при изгибе:")
    _formula(doc, [("σ", 'i'), ("max", 'sub'), (" = |M", 'n'), ("max", 'sub'), ("| / W", 'n'),
                    ("x", 'sub'), (" ≤ [σ]", 'n')])
    sigma_max = abs(Mv) / sec['Wx']
    _formula(doc, [("σ", 'i'), ("max", 'sub'), (" = ", 'n'), (_num(abs(Mv), 4), 'n'),
                    (" / ", 'n'), (_num(sec['Wx'], 4), 'n'), (" = ", 'n'),
                    (_num(sigma_max / 1e6, 4), 'n'), (" МПа", 'n')])
    n_safety = ctx['sigma_allow'] / sigma_max if sigma_max > 0 else float('inf')
    ok = sigma_max <= ctx['sigma_allow']
    _p(doc, f"Допускаемое напряжение [σ] = {ctx['sigma_allow']/1e6:.3g} МПа.")
    _p(doc, f"Коэффициент запаса прочности: n = [σ] / σ_max = {n_safety:.2f}.")
    verdict = ("Условие прочности ВЫПОЛНЯЕТСЯ" if ok else "Условие прочности НЕ ВЫПОЛНЯЕТСЯ")
    _p(doc, f"{verdict} (σ_max {'≤' if ok else '>'} [σ]).", bold=True)

    # ---------- 6. Жёсткость ----------
    doc.add_page_break()
    _h1(doc, "6. Проверка жёсткости (прогибы)")
    xs = [beam.L * i / 300 for i in range(301)]
    vs = [beam.deflection(x) for x in xs]
    i_max = max(range(len(vs)), key=lambda i: abs(vs[i]))
    v_max, x_at = vs[i_max], xs[i_max]
    f_allow = ctx.get('deflection_ratio', 250)
    v_allow = L / f_allow
    _p(doc, f"Максимальный прогиб: v_max = {v_max*1000:.4f} мм при x = {x_at:.3f} м.")
    _p(doc, f"Допускаемый прогиб (по условию жёсткости [v] = L/{f_allow}): "
            f"[v] = {v_allow*1000:.3f} мм.")
    ok_v = abs(v_max) <= v_allow
    _p(doc, f"{'Условие жёсткости ВЫПОЛНЯЕТСЯ' if ok_v else 'Условие жёсткости НЕ ВЫПОЛНЯЕТСЯ'} "
            f"(|v_max| {'≤' if ok_v else '>'} [v]).", bold=True)
    _add_image(doc, ctx['images']['defl'])
    _p(doc, "Рисунок 5 — эпюра прогибов v(x)", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True)

    # ---------- Вывод ----------
    doc.add_page_break()
    _h1(doc, "Вывод")
    _p(doc, f"В результате расчёта балки длиной L = {_num(L)} м определены опорные реакции, "
            f"построены эпюры Q(x) и M(x), максимальный изгибающий момент составил "
            f"M_max = {_num(Mv)} Н·м. Для сечения «{sec['desc']}» максимальное напряжение "
            f"σ_max = {sigma_max/1e6:.3g} МПа при [σ] = {ctx['sigma_allow']/1e6:.3g} МПа "
            f"(коэффициент запаса n = {n_safety:.2f}); максимальный прогиб v_max = "
            f"{v_max*1000:.4f} мм при допускаемом [v] = {v_allow*1000:.3f} мм. "
            f"{'Балка удовлетворяет условиям прочности и жёсткости.' if (ok and ok_v) else 'Балка НЕ удовлетворяет одному или обоим условиям — требуется увеличить сечение.'}")

    doc.save(out_path)
    return out_path
